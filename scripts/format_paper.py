#!/usr/bin/env python3
"""
武汉大学国家网络安全学院 课程论文/实验报告 格式套用工具

按学院排版规范对论文或实验报告内容进行格式化，输出符合要求的 .docx 文件。
不包含封面生成——仅负责内容排版。

支持两种文档类型：
  --type paper   课程论文（默认）
  --type report  实验报告
"""

import json
import re
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── 字号映射 ──────────────────────────────────────────

FONT_SIZES = {
    "2号": Pt(22),
    "小2号": Pt(18),
    "4号": Pt(14),
    "小4号": Pt(12),
    "5号": Pt(10.5),
}


# ── 页面设置 ──────────────────────────────────────────

def setup_page(doc):
    """A4纸，页边距：上25mm 下20mm 左30mm 右30mm。"""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)


def add_page_number(doc):
    """页脚居中页码，Times New Roman 5号。"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        run2 = p.add_run()
        run2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        run3 = p.add_run()
        run3._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

        for r in [run, run2, run3]:
            r.font.size = FONT_SIZES["5号"]
            r.font.name = "Times New Roman"


# ── 样式定义 ──────────────────────────────────────────

def define_styles(doc):
    """定义正文和标题样式。"""
    style = doc.styles

    normal = style["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = FONT_SIZES["小4号"]
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = normal.paragraph_format
    pf.line_spacing = Pt(23)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    _set_heading(style, "Heading 1", "黑体", FONT_SIZES["小2号"], Pt(9.6), Pt(6))
    _set_heading(style, "Heading 2", "黑体", FONT_SIZES["4号"], Pt(6), Pt(6))
    _set_heading(style, "Heading 3", "黑体", FONT_SIZES["小4号"], Pt(6), Pt(6))
    _set_heading(style, "Heading 4", "黑体", FONT_SIZES["小4号"], Pt(6), Pt(6))


def _set_heading(styles, name, font_name, font_size, space_before, space_after):
    try:
        s = styles[name]
    except KeyError:
        return
    s.font.name = font_name
    s.font.size = font_size
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    s.paragraph_format.space_before = space_before
    s.paragraph_format.space_after = space_after
    s.paragraph_format.line_spacing = Pt(23)
    s.paragraph_format.keep_with_next = True


# ── 元素写入 ──────────────────────────────────────────

def add_section_heading(doc, title, level):
    """添加章节标题（黑体，层级对应不同字号）。"""
    p = doc.add_paragraph()
    run = p.add_run(title)

    sizes = {1: FONT_SIZES["小2号"], 2: FONT_SIZES["4号"], 3: FONT_SIZES["小4号"]}
    spacing_before = {1: Pt(9.6), 2: Pt(6), 3: Pt(6)}
    spacing_after = {1: Pt(6), 2: Pt(6), 3: Pt(6)}

    run.font.size = sizes.get(level, FONT_SIZES["小4号"])
    run.font.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p.paragraph_format.space_before = spacing_before.get(level, Pt(6))
    p.paragraph_format.space_after = spacing_after.get(level, Pt(6))
    p.paragraph_format.line_spacing = Pt(23)
    return p


def add_body_paragraph(doc, text):
    """正文段落：宋体小4号，固定行距23磅，首行缩进2字符。"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(23)
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = FONT_SIZES["小4号"]
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_abstract(doc, abstract_text, keywords, lang="cn"):
    """添加中/英文摘要与关键词。英文摘要另起一页。"""
    if lang == "en":
        doc.add_page_break()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ABSTRACT")
        run.font.size = FONT_SIZES["小2号"]
        run.font.bold = True
        run.font.name = "Times New Roman"

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(23)
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(abstract_text)
        run.font.size = FONT_SIZES["小4号"]
        run.font.name = "Times New Roman"

        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("Keywords: ")
        run.font.size = FONT_SIZES["小4号"]
        run.font.bold = True
        run.font.name = "Times New Roman"
        run = p.add_run("; ".join(keywords))
        run.font.size = FONT_SIZES["小4号"]
        run.font.name = "Times New Roman"
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("摘  要")
        run.font.size = FONT_SIZES["小2号"]
        run.font.bold = True
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        p.paragraph_format.space_after = Pt(6)

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(23)
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(abstract_text)
        run.font.size = FONT_SIZES["小4号"]
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        doc.add_paragraph()

        p = doc.add_paragraph()
        run = p.add_run("关键词：")
        run.font.size = FONT_SIZES["小4号"]
        run.font.bold = True
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run = p.add_run("；".join(keywords))
        run.font.size = FONT_SIZES["小4号"]
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_references(doc, references):
    """参考文献列表，标题黑体小2号居中，条目宋体小4号。

    自动去除输入中已有的 [N] 前缀并重新编号。
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("参考文献")
    run.font.size = FONT_SIZES["小2号"]
    run.font.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p.paragraph_format.space_after = Pt(6)

    for i, ref in enumerate(references, 1):
        ref_clean = re.sub(r'^\[\d+\]\s*', '', ref.strip())
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(23)
        run = p.add_run(f"[{i}] {ref_clean}")
        run.font.size = FONT_SIZES["小4号"]
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


# ── 文档结构渲染 ──────────────────────────────────────

def _render_sections(doc, sections):
    """将 sections 列表渲染为标题 + 正文段落。"""
    for sec in sections:
        add_section_heading(doc, sec.get("title", ""), sec.get("level", 1))
        for txt in sec.get("content", []):
            add_body_paragraph(doc, txt)
        for sub in sec.get("subsections", []):
            add_section_heading(doc, sub.get("title", ""), sub.get("level", 2))
            for txt in sub.get("content", []):
                add_body_paragraph(doc, txt)


def _render_document(doc, content):
    """渲染论文公共部分：摘要（默认仅中文）、章节、参考文献。"""
    # 中文摘要（默认生成）
    if content.get("abstract_cn"):
        add_abstract(doc, content["abstract_cn"], content.get("keywords_cn", []), lang="cn")
    # 英文摘要（仅在用户明确要求时生成）
    if content.get("abstract_en"):
        add_abstract(doc, content["abstract_en"], content.get("keywords_en", []), lang="en")
    # 章节
    sections = content.get("sections", [])
    if sections:
        _render_sections(doc, sections)
    # 参考文献
    refs = content.get("references", [])
    if refs:
        add_references(doc, refs)


# ── 论文入口 ──────────────────────────────────────────

def format_paper(output_path, content):
    """格式化课程论文。"""
    doc = _create_document()
    _render_document(doc, content)
    doc.save(output_path)
    print(f"论文已格式化: {output_path}")
    return output_path


# ── 报告入口 ──────────────────────────────────────────

def format_report(output_path, content):
    """格式化实验报告。

    content 结构:
    {
        "title": "实验报告标题",         // 可选，报告总标题（黑体二号居中）
        "purpose": {"title": "实验目的", "content": ["..."]},
        "environment": {"title": "实验环境", "content": ["..."]},
        "experiment_content": {"title": "实验内容", "content": ["..."]},
        "steps": {"title": "实验步骤", "content": ["..."]},
        "results": {"title": "实验结果与分析", "content": ["..."]},
        "summary": {"title": "实验总结", "content": ["..."]}
    }
    """
    doc = _create_document()

    # 报告标题（可选，居中加粗）
    title = content.get("title", "")
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = FONT_SIZES["2号"]
        run.font.bold = True
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        p.paragraph_format.space_after = Pt(12)

    # 报告固定模块：按顺序渲染
    REPORT_SECTIONS = [
        ("purpose", "实验目的"),
        ("environment", "实验环境"),
        ("experiment_content", "实验内容"),
        ("steps", "实验步骤"),
        ("results", "实验结果与分析"),
        ("summary", "实验总结"),
    ]

    for key, default_title in REPORT_SECTIONS:
        sec = content.get(key)
        if sec is None:
            continue
        title_text = sec.get("title", default_title)
        body_texts = sec.get("content", [])

        add_section_heading(doc, title_text, level=1)
        for txt in body_texts:
            add_body_paragraph(doc, txt)

    doc.save(output_path)
    print(f"报告已格式化: {output_path}")
    return output_path


# ── 文档初始化 ──────────────────────────────────────────

def _create_document():
    """创建并初始化一个符合规范的空白文档。"""
    doc = Document()
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    setup_page(doc)
    define_styles(doc)
    add_page_number(doc)
    return doc


# ── CLI ────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="武汉大学国家网络安全学院论文/报告格式套用")
    parser.add_argument("--output", "-o", required=True, help="输出 .docx 路径")
    parser.add_argument("--content", "-c", required=True, help="内容 JSON 文件路径")
    parser.add_argument("--type", "-t", choices=["paper", "report"], default="paper",
                        help="文档类型：paper（课程论文）或 report（实验报告）")
    args = parser.parse_args()

    with open(args.content, "r", encoding="utf-8") as f:
        content = json.load(f)

    if args.type == "report":
        format_report(args.output, content)
    else:
        format_paper(args.output, content)
